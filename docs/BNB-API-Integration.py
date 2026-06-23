"""
Genera el documento Word "Integracion API BNB - Festival Danzarte 2026".
Ejecutar: python BNB-API-Integration.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_code(doc, code, lang=""):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F1F5F9')
    p_pr.append(shd)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)


def add_endpoint(doc, num, title, method, path, sandbox, descr, params, req_example, resp_fields, notes=None):
    doc.add_heading(f"{num}. {title}", level=2)

    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    cells = table.rows[0].cells
    cells[0].text = "Método"
    cells[1].text = method
    cells = table.rows[1].cells
    cells[0].text = "URL"
    cells[1].text = path
    cells = table.rows[2].cells
    cells[0].text = "Sandbox"
    cells[1].text = sandbox
    cells = table.rows[3].cells
    cells[0].text = "Descripción"
    cells[1].text = descr

    doc.add_paragraph()
    doc.add_paragraph("Parámetros:", style='Heading 3').paragraph_format.space_before = Pt(8)
    if params:
        tbl = doc.add_table(rows=1 + len(params), cols=3)
        tbl.style = 'Light List Accent 1'
        hdr = tbl.rows[0].cells
        hdr[0].text = "Campo"
        hdr[1].text = "Tipo"
        hdr[2].text = "Descripción"
        for i, (n, t, d) in enumerate(params, start=1):
            row = tbl.rows[i].cells
            row[0].text = n
            row[1].text = t
            row[2].text = d

    doc.add_paragraph()
    doc.add_paragraph("Request:", style='Heading 3').paragraph_format.space_before = Pt(8)
    add_code(doc, req_example)

    doc.add_paragraph("Respuesta:", style='Heading 3').paragraph_format.space_before = Pt(8)
    if isinstance(resp_fields, str):
        add_code(doc, resp_fields)
    else:
        tbl = doc.add_table(rows=1 + len(resp_fields), cols=2)
        tbl.style = 'Light List Accent 1'
        hdr = tbl.rows[0].cells
        hdr[0].text = "Campo"
        hdr[1].text = "Significado"
        for i, (n, d) in enumerate(resp_fields, start=1):
            row = tbl.rows[i].cells
            row[0].text = n
            row[1].text = d

    if notes:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Notas para festival: ")
        run.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0xB6, 0xD9)
        p.add_run(notes)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)


def main():
    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.2)

    # Default font Helvetica/Inter
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ───────────────── PORTADA ─────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("INTEGRACIÓN BNB OPEN BANKING API")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x0A, 0x40)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Festival Danzarte 2026 — Análisis y guía técnica")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Versión 1.0 · 15 de mayo de 2026\n")
    info.add_run("Banco Nacional de Bolivia · Open Banking API v2.0")

    doc.add_page_break()

    # ───────────────── 1. RESUMEN ─────────────────
    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "El Banco Nacional de Bolivia expone una API REST que permite generar códigos QR "
        "de cobro, consultar su estado, recibir transferencias y consultar saldos de cuenta. "
        "El presente documento detalla los endpoints disponibles, identifica los que "
        "necesitamos para el Festival Danzarte 2026 y describe la arquitectura de "
        "notificaciones que conectaremos a nuestro flujo n8n para acreditar pagos "
        "automáticamente en la app del portal de usuarios."
    )

    doc.add_heading("Arquitectura propuesta", level=2)
    doc.add_paragraph(
        "1. El usuario abre el modal de pago en la app y pulsa “Generar QR de pago”.\n"
        "2. El backend PHP solicita un token a BNB y luego un QR de monto fijo "
        "vinculado a la inscripción.\n"
        "3. La app muestra el QR; el usuario lo escanea desde su app bancaria.\n"
        "4. Cuando BNB acredita el pago, el QR pasa al estado “2 = Used”. n8n hace "
        "polling cada 30 segundos contra el endpoint getQRStatusAsync y, al detectar "
        "el cambio, marca el pago como “verificado” en Supabase, genera el PDF de "
        "recibo y envía la notificación al usuario."
    )

    doc.add_heading("Endpoints que vamos a usar", level=2)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = tbl.rows[0].cells
    hdr[0].text = "Uso"
    hdr[1].text = "Endpoint"
    hdr[2].text = "Frecuencia"
    rows_data = [
        ("Autenticación", "POST /auth/token", "Cada 15 min (token cache)"),
        ("Generar QR de cobro", "POST /main/getQRWithImageAsync", "Por cada pago iniciado"),
        ("Consultar estado QR", "POST /main/getQRStatusAsync", "Polling 30 s (n8n)"),
        ("Listar QRs por fecha", "POST /main/getQRbyGenerationDateAsync", "Reconciliación diaria"),
        ("Saldo de cuenta", "POST /Transactions/getBalanceAsync", "Dashboard admin"),
    ]
    for u, e, f in rows_data:
        row = tbl.add_row().cells
        row[0].text = u
        row[1].text = e
        row[2].text = f

    doc.add_page_break()

    # ───────────────── 2. AUTENTICACIÓN ─────────────────
    doc.add_heading("2. Autenticación", level=1)
    doc.add_paragraph(
        "Todos los endpoints de QR Simple y Direct Debit requieren un token Bearer "
        "que se obtiene con el endpoint de autenticación. El token tiene una "
        "vigencia limitada (verificar con BNB; típicamente 15-60 min). Recomendamos "
        "cachearlo en el backend PHP y renovarlo bajo demanda."
    )
    add_endpoint(
        doc, "2.1", "Generación de token",
        "POST",
        "/ClientAuthentication.API/api/v1/auth/token",
        "http://test.bnb.com.bo/ClientAuthentication.API/api/v1/auth/token",
        "Devuelve un token de seguridad para autorizar las llamadas a los demás endpoints.",
        [
            ("accountId", "string", "Identificador de cuenta (provisto por BNB al alta)"),
            ("authorizationId", "string", "Identificador de autorización (provisto por BNB)"),
        ],
        '{\n  "accountId": "__BNB_ACCOUNT_ID__",\n  "authorizationId": "__BNB_AUTH_ID__"\n}',
        [
            ("success", "true si la autenticación fue exitosa"),
            ("message", "Token Bearer (cuando success=true) o mensaje de error"),
        ],
        "Almacenar el token en cache de Redis o en archivo temp con TTL de 10 min. "
        "Renovar automáticamente cuando una llamada devuelva 401."
    )

    doc.add_page_break()

    # ───────────────── 3. QR DE COBRO ─────────────────
    doc.add_heading("3. Generación de QR de cobro (caso principal)", level=1)
    doc.add_paragraph(
        "Este es el flujo central para el Festival. Cuando un representante decide "
        "pagar la inscripción, kardex o credenciales, generamos un QR con monto, "
        "vencimiento y referencia única vinculada al pago en nuestra base."
    )

    add_endpoint(
        doc, "3.1", "QR Simple con Imagen",
        "POST",
        "/main/getQRWithImageAsync",
        "https://qrsimpleapiv2.azurewebsites.net/api/v1/",
        "Genera un QR de cobro y lo devuelve como imagen (byte array). El QR puede ser "
        "de uso único o reutilizable.",
        [
            ("currency", "string", "Moneda: BOB o USD"),
            ("gloss", "string", "Descripción del cobro (aparece en el extracto)"),
            ("amount", "string", "Monto a cobrar"),
            ("expirationDate", "string", "Fecha de expiración del QR (yyyy-MM-dd)"),
            ("singleUse", "boolean", "true = un solo uso; false = reutilizable"),
        ],
        '{\n  "currency": "BOB",\n  "gloss": "Inscripción SHERA PRUEBA · Festival Danzarte 2026",\n  "amount": "700",\n  "expirationDate": "2026-06-29",\n  "singleUse": true\n}',
        [
            ("id", "Identificador único del QR generado (lo guardamos en pagos_2026.qr_id)"),
            ("qr", "Imagen del QR en bytes (base64) — la mostramos en la app"),
            ("success", "true si se generó correctamente"),
            ("message", "Mensaje de éxito o detalle del error"),
        ],
        "Persistir el campo id en la columna pagos_2026.qr_id para poder consultar "
        "el estado más tarde. La gloss debe incluir el id de inscripción para que sea "
        "fácil reconciliar manualmente si fuera necesario."
    )

    add_endpoint(
        doc, "3.2", "Consultar estado del QR",
        "POST",
        "/main/getQRStatusAsync",
        "https://qrsimpleapiv2.azurewebsites.net/api/v1/",
        "Devuelve el estado actual del QR. Este endpoint es la base del polling que "
        "hará n8n para detectar pagos acreditados.",
        [
            ("qrId", "string", "Identificador del QR (id devuelto por getQRWithImageAsync)"),
        ],
        '{\n  "qrId": "59"\n}',
        [
            ("id", "Identificador del QR"),
            ("qrId", "Estado (entero): 1 = No usado · 2 = Usado · 3 = Expirado · 4 = Error"),
            ("expirationDate", "Fecha de expiración"),
            ("success", "true / false"),
            ("message", "Mensaje"),
        ],
        "Cuando el estado pase a 2 (Usado), n8n debe marcar pagos_2026.estado = "
        "'verificado', actualizar verificado_en y disparar la generación del recibo PDF."
    )

    add_endpoint(
        doc, "3.3", "Listar QRs por fecha de generación",
        "POST",
        "/main/getQRbyGenerationDateAsync",
        "https://qrsimpleapiv2.azurewebsites.net/api/v1/",
        "Devuelve los QRs generados en una fecha específica. Útil para la "
        "reconciliación diaria al final del día.",
        [
            ("generationDate", "string", "Fecha de generación (yyyy-MM-dd)"),
        ],
        '{\n  "generationDate": "2026-06-28"\n}',
        [
            ("dTOqrDetails", "Lista con detalles de cada QR del día"),
            ("success", "true / false"),
            ("message", "Mensaje"),
        ],
        "Configurar un workflow n8n diario a las 23:55 que llame este endpoint, "
        "compare con pagos_2026 y notifique al admin si hay diferencias."
    )

    doc.add_page_break()

    # ───────────────── 4. CONSULTA DE CUENTA ─────────────────
    doc.add_heading("4. Consultas de cuenta y movimientos", level=1)
    doc.add_paragraph(
        "Endpoints útiles para el panel administrativo del festival: ver saldo "
        "de la cuenta receptora en tiempo real y los últimos movimientos."
    )

    add_endpoint(
        doc, "4.1", "Saldo de la cuenta autenticada",
        "POST",
        "/Transactions/getBalanceAsync",
        "https://accountapiv1.azurewebsites.net/api/v1/",
        "Devuelve el saldo de todas las cuentas vinculadas al cliente "
        "autenticado. No requiere parámetros adicionales en el body.",
        [],
        '{}',
        [
            ("accountBalance[].accountNumber", "Número de cuenta"),
            ("accountBalance[].accountType", "1 = Corriente · 2 = Caja de ahorro"),
            ("accountBalance[].Currency", "Moneda de la cuenta"),
            ("accountBalance[].balanceAmount", "Saldo disponible"),
            ("partyName", "Razón social del titular"),
        ],
        "Mostrar el saldo en el dashboard admin con refresh cada 5 minutos."
    )

    add_endpoint(
        doc, "4.2", "Saldo de una cuenta específica",
        "POST",
        "/Transactions/getAccountBalanceAsync",
        "https://accountapiv1.azurewebsites.net/api/v1/",
        "Devuelve el saldo de una cuenta específica del cliente.",
        [
            ("AccountNumber", "string", "Número de cuenta a consultar"),
        ],
        '{\n  "AccountNumber": "1501243627"\n}',
        [
            ("accountBalance", "Información de la cuenta consultada"),
            ("partyName", "Razón social"),
        ],
        None
    )

    add_endpoint(
        doc, "4.3", "Últimos 10 movimientos (Bank Statement)",
        "POST",
        "/Enterprise/BankStatement",
        "http://bnbapideveloperv1.azurewebsites.net/",
        "Devuelve los últimos 10 movimientos de la cuenta.",
        [
            ("userKey", "string", "Llave del usuario empresarial"),
            ("accountNumber", "string", "Número de cuenta"),
        ],
        '{\n  "userKey": "e8k7crKA9S0:APA91bGDZ76NccQkYXIzS5",\n  "accountNumber": "1501243627"\n}',
        [
            ("accountNumber", "Cuenta consultada"),
            ("accountType", "Tipo de cuenta"),
            ("movements", "Lista de transacciones"),
        ],
        None
    )

    doc.add_page_break()

    # ───────────────── 5. DIRECT DEBIT ─────────────────
    doc.add_heading("5. Débito automático (Direct Debit) — uso opcional", level=1)
    doc.add_paragraph(
        "Estos endpoints permiten configurar débitos automáticos. Para el festival "
        "normalmente NO se usan (los participantes pagan QR puntuales), pero los "
        "documentamos por si en el futuro se ofrece un convenio de pago en cuotas."
    )

    add_endpoint(
        doc, "5.1", "QR de monto fijo (Direct Debit)",
        "POST",
        "/Services/GetQRFixedAmount",
        "http://test.bnb.com.bo/DirectDebit/api/Services/GetQRFixedAmount",
        "Genera un QR de domiciliación con monto fijo y cuotas predefinidas.",
        [
            ("currencyCode", "integer", "1 = BOB · 2 = USD"),
            ("amount", "number", "Monto fijo (> 0, hasta 2 decimales)"),
            ("reference", "string", "Descripción del servicio"),
            ("serviceCode", "string", "Código de servicio / contrato"),
            ("dueDate", "string", "Fecha de expiración (yyyy-MM-dd)"),
            ("installmentsQuantity", "integer", "Cantidad de cuotas"),
            ("chargeType", "integer", "1 = automático · 2 = manual"),
            ("chargeDate", "integer", "Día de cobro (sólo modo automático)"),
        ],
        '{\n  "currencyCode": 1,\n  "amount": 700,\n  "reference": "Inscripción Festival Danzarte",\n  "serviceCode": "FDZ2026",\n  "dueDate": "2026-06-29",\n  "installmentsQuantity": 1,\n  "chargeType": 2\n}',
        [
            ("method", '"QR"'),
            ("qrId", "Código único de domiciliación"),
            ("qrContent", "Imagen base64 del QR"),
            ("installments[]", "Plan de cuotas"),
            ("success / message / code", "Estado"),
        ],
        None
    )

    add_endpoint(
        doc, "5.2", "QR de monto variable",
        "POST",
        "/Services/GetQRVariableAmount",
        "http://test.bnb.com.bo/DirectDebit/api/Services/GetQRVariableAmount",
        "Genera un QR de domiciliación de monto variable (válido hasta revocar).",
        [("(idénticos a 5.1)", "—", "Mismos campos")],
        '{\n  "currencyCode": 1,\n  "amount": 0,\n  "reference": "Mensualidad academia",\n  "serviceCode": "FDZ-VAR-2026",\n  "dueDate": "2026-12-31",\n  "installmentsQuantity": 12,\n  "chargeType": 1,\n  "chargeDate": 15\n}',
        "(Mismos campos de respuesta que 5.1)",
        None
    )

    other_dd = [
        ("5.3", "Enviar orden de débito", "/Services/SendDebitOrder", "Solicita el débito de una cuota específica."),
        ("5.4", "Listar domiciliaciones generadas", "/Services/GetTransactionOutgoing", "Lista las domiciliaciones en un rango (máx. 30 días)."),
        ("5.5", "Cobros rechazados", "/Services/GetTransactionRejected", "Cobros que fueron rechazados."),
        ("5.6", "Detalle de domiciliación", "/Services/GetDetail", "Obtiene los detalles completos a partir del qrId."),
        ("5.7", "Actualizar cuota pendiente", "/Services/UpdatePendingQuota", "Modifica el monto/fecha de una cuota."),
    ]
    for n, t, p, d in other_dd:
        doc.add_heading(f"{n}. {t}", level=2)
        doc.add_paragraph(f"POST {p}")
        doc.add_paragraph(d).paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ───────────────── 6. WEBHOOKS / NOTIFICACIONES ─────────────────
    doc.add_heading("6. Notificaciones de pago hacia n8n", level=1)

    doc.add_paragraph(
        "Importante: BNB no expone públicamente un webhook estilo Stripe en el "
        "sandbox actual. Las notificaciones de pago se obtienen mediante polling "
        "del endpoint getQRStatusAsync (sección 3.2). El patrón recomendado es:"
    )

    doc.add_heading("Patrón A — Polling con n8n (recomendado)", level=2)
    doc.add_paragraph(
        "Crear un workflow en n8n con un trigger Cron cada 30 segundos. El flujo:"
    )
    steps = [
        "Trigger Cron (intervalo 30 s).",
        "Nodo HTTP a Supabase: SELECT qr_id FROM pagos_2026 WHERE estado='enviado' AND qr_id IS NOT NULL.",
        "Loop sobre cada qr_id pendiente.",
        "Nodo HTTP POST a /auth/token (con cache de 10 min en n8n Variables).",
        "Nodo HTTP POST a /main/getQRStatusAsync con el qrId.",
        "Si response.qrId == 2 (Usado): actualizar pagos_2026 SET estado='verificado', verificado_en=now() vía RPC Supabase.",
        "Disparar la generación del PDF llamando a POST /recibo-generar.php.",
        "Enviar notificación push/email al usuario.",
        "Si response.qrId == 3 (Expirado): marcar pagos_2026.estado='anulado'.",
    ]
    for i, s in enumerate(steps, start=1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(s)

    doc.add_heading("Workflow n8n — JSON template", level=2)
    add_code(doc, """{
  "name": "BNB QR Status Poller",
  "nodes": [
    { "name": "Cron 30s",   "type": "n8n-nodes-base.cron",      "parameters": { "triggerTimes": { "item": [{ "mode": "everyX", "value": 30, "unit": "seconds" }] } } },
    { "name": "Pagos pendientes", "type": "n8n-nodes-base.httpRequest",
      "parameters": {
        "url": "https://supabase.imaginarte.cloud/rest/v1/pagos_2026?select=id_pago,qr_id&estado=eq.enviado&qr_id=not.is.null",
        "options": { "headers": { "apikey": "={{$env.SUPABASE_SERVICE_KEY}}", "Authorization": "=Bearer {{$env.SUPABASE_SERVICE_KEY}}" } }
      }
    },
    { "name": "Split In Batches", "type": "n8n-nodes-base.splitInBatches" },
    { "name": "BNB Token",        "type": "n8n-nodes-base.httpRequest", "parameters": { "method": "POST", "url": "https://clientauthenticationapiv2.azurewebsites.net/api/v1/auth/token", "jsonParameters": true, "bodyParametersJson": "={ \\"accountId\\": \\"{{$env.BNB_ACCOUNT_ID}}\\", \\"authorizationId\\": \\"{{$env.BNB_AUTH_ID}}\\" }" } },
    { "name": "Get QR Status",    "type": "n8n-nodes-base.httpRequest", "parameters": { "method": "POST", "url": "https://qrsimpleapiv2.azurewebsites.net/api/v1/main/getQRStatusAsync", "options": { "headers": { "Authorization": "=Bearer {{$node[\\"BNB Token\\"].json[\\"message\\"]}}" } }, "jsonParameters": true, "bodyParametersJson": "={ \\"qrId\\": \\"{{$json.qr_id}}\\" }" } },
    { "name": "IF Usado",         "type": "n8n-nodes-base.if", "parameters": { "conditions": { "number": [{ "value1": "={{$json.qrId}}", "operation": "equal", "value2": 2 }] } } },
    { "name": "Marcar verificado","type": "n8n-nodes-base.httpRequest", "parameters": { "method": "PATCH", "url": "=https://supabase.imaginarte.cloud/rest/v1/pagos_2026?id_pago=eq.{{$json.id_pago}}", "options": { "headers": { "apikey": "={{$env.SUPABASE_SERVICE_KEY}}", "Authorization": "=Bearer {{$env.SUPABASE_SERVICE_KEY}}", "Content-Type": "application/json", "Prefer": "return=minimal" } }, "jsonParameters": true, "bodyParametersJson": "={ \\"estado\\": \\"verificado\\", \\"verificado_en\\": \\"{{$now.toISO()}}\\", \\"verificado_por\\": \\"BNB_AUTO\\" }" } },
    { "name": "Generar recibo",   "type": "n8n-nodes-base.httpRequest", "parameters": { "method": "POST", "url": "https://festivaldanzarte.com/api/recibo-generar.php", "jsonParameters": true, "bodyParametersJson": "={ \\"id_pago\\": \\"{{$json.id_pago}}\\" }" } }
  ]
}""")

    doc.add_heading("Patrón B — Webhook directo BNB (si lo habilitan)", level=2)
    doc.add_paragraph(
        "BNB Commerce menciona “alertas de cobro” en su material comercial, pero el "
        "endpoint público de webhook no está documentado en el portal Sandbox. Pasos:"
    )
    bsteps = [
        "Contactar al equipo comercial de BNB y solicitar habilitación de webhook saliente para nuestro accountId.",
        "Configurar como callback URL: https://danzarte-n8n.rgxmhp.easypanel.host/webhook/bnb-pago-recibido",
        "BNB hará POST con el payload (estructura típica: { qrId, amount, currency, transactionId, paidAt }).",
        "n8n recibe el webhook → valida HMAC firma (si BNB lo provee) → actualiza pagos_2026 igual que en el patrón A.",
    ]
    for i, s in enumerate(bsteps, start=1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(s)

    add_code(doc, '{\n  "qrId": "59",\n  "amount": 700,\n  "currency": "BOB",\n  "transactionId": "TXN-20260629-001",\n  "paidAt": "2026-06-29T14:32:11-04:00",\n  "payerName": "JUAN PÉREZ",\n  "payerAccount": "*****4321"\n}')

    doc.add_page_break()

    # ───────────────── 7. CHECKLIST DE INTEGRACIÓN ─────────────────
    doc.add_heading("7. Checklist de integración", level=1)
    items = [
        "Solicitar a BNB credenciales productivas (accountId, authorizationId, userKey).",
        "Agregar variables a php-backend/config.php: bnb_account_id, bnb_authorization_id, bnb_user_key, bnb_token_cache_ttl.",
        "Crear php-backend/_lib/bnb.php con funciones bnbToken(), bnbGenerarQR(monto, gloss, ref), bnbConsultarQR(qrId).",
        "Modificar pago-crear.php para que tras crear la fila pagos_2026 invoque bnbGenerarQR() y guarde qr_id, qr_url.",
        "Cambiar el botón \"Generar QR\" del PagoModal para que muestre el QR generado dinámicamente (no la imagen estática).",
        "Crear workflow n8n descrito en sección 6 y configurar variables de entorno con las credenciales.",
        "Agregar columna pagos_2026.qr_id (text) si aún no existe.",
        "Pruebas en sandbox con cuenta de prueba antes de habilitar producción.",
        "Documentar en el README del proyecto el flujo y las credenciales requeridas.",
    ]
    for s in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(s)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Próximo paso: cuando reciba las credenciales productivas, "
                  "actualice php-backend/config.php y avísenme para implementar "
                  "bnb.php y modificar pago-crear.php.")
    r.italic = True
    r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Save
    out = r"D:\Claude\APPS\APP FESTIVAL DANZARTE SISTEMA DE USUARIOS - VITE\docs\BNB-API-Integration-Festival-Danzarte.docx"
    doc.save(out)
    print(f"OK: {out}")


if __name__ == "__main__":
    main()
